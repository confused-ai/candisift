"""Text-extraction adapter: PDF / DOCX / image / plain text -> string.

Implements the TextExtractor port. Strategy per type:

- **PDF**: read the embedded text layer (pdfplumber). If a page has little/no
  text (scanned image PDF), render that page and OCR it with Tesseract. Pages
  that already have a text layer are kept verbatim — OCR runs only where needed.
- **Image** (png/jpg/tiff/…): OCR directly with Tesseract.
- **DOCX**: read paragraphs and table cells (python-docx); embedded images are
  OCR'd when OCR is on.
- **Legacy .doc** (Word 97-2003, incl. one renamed to .docx): converted first
  via `textutil` (stock on macOS) then `soffice`, and read as text/DOCX.
- **Everything else**: decode as UTF-8 (lossy).

Routing is by magic bytes first, extension only as a fallback — mislabeled
uploads (a .doc or PDF named .docx) are common and used to yield no text at all.

Every optional dependency (pdfplumber, python-docx, pytesseract, pdf2image,
Pillow) and the external `tesseract`/`poppler` binaries are imported/called
lazily inside try/except. A missing dep or a corrupt file degrades to the best
text we already have (or "") and logs a warning — ingestion never crashes.

System deps for OCR: `tesseract` (the OCR engine) and `poppler` (PDF->image,
used by pdf2image). macOS: `brew install tesseract poppler`. Debian/Ubuntu:
`apt-get install tesseract-ocr poppler-utils`.
"""
from __future__ import annotations

import io
import logging
import os

log = logging.getLogger("candisift.parsing")

# extensions handled as raster images (OCR-only)
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp", ".gif", ".ppm", ".pnm"}

# a page/image yielding fewer than this many non-space chars is treated as
# "no real text layer" -> OCR it. Tuned to skip headers/page numbers only.
_OCR_MIN_CHARS = 24

# decompression-bomb guard: refuse to rasterize an image above this many pixels.
# A 300-DPI A4 page is ~8.7MP; 40MP leaves headroom while blocking a crafted file
# that expands to gigapixels and exhausts memory. Pillow raises on exceeding it.
_MAX_IMAGE_PIXELS = 40_000_000

# OLE2 / Compound File Binary header — legacy Word 97-2003 (.doc), and what a
# .doc renamed to .docx actually contains.
_OLE2_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"


class FileTextExtractor:
    """Filename-dispatched extractor with an optional Tesseract OCR fallback.

    Args:
        ocr: master switch. When False, behaves as a pure text-layer extractor.
        ocr_lang: Tesseract language(s), e.g. "eng" or "eng+deu".
        ocr_dpi: render resolution for PDF pages. 300 is the OCR sweet spot;
            higher is slower with little accuracy gain on resumes.
        max_ocr_pages: hard cap on pages rendered+OCR'd per PDF (memory/time
            guard against a 500-page scan). Pages beyond it keep their text
            layer if any, else are skipped.
    """

    def __init__(
        self,
        ocr: bool = True,
        ocr_lang: str = "eng",
        ocr_dpi: int = 300,
        max_ocr_pages: int = 50,
        ocr_timeout_s: int = 30,
        doc_convert_timeout_s: int = 60,
    ) -> None:
        self.ocr = ocr
        # LibreOffice cold-starts slowly on the first .doc of a process lifetime
        self.doc_convert_timeout_s = max(5, int(doc_convert_timeout_s))
        self.ocr_lang = ocr_lang
        self.ocr_dpi = max(72, int(ocr_dpi))
        self.max_ocr_pages = max(1, int(max_ocr_pages))
        # per-call wall-clock cap on the tesseract/poppler subprocesses so a
        # hostile or pathological file can't wedge a worker thread indefinitely.
        self.ocr_timeout_s = max(1, int(ocr_timeout_s))

    # ---- port entrypoint -------------------------------------------------
    def extract(self, content: bytes, filename: str) -> str:
        if not content:
            return ""
        ext = os.path.splitext((filename or "").lower())[1]
        # Content wins over extension. A legacy .doc (or a PDF) renamed to .docx
        # is routine from recruiters/applicants, and used to die as BadZipFile
        # with zero text. Magic bytes are unambiguous; the extension is only a
        # hint for the formats we can't sniff (.txt/.md/.csv).
        if content.startswith(b"%PDF"):
            return self._pdf(content)
        if content.startswith(b"PK\x03\x04"):
            return self._docx(content)
        if content.startswith(_OLE2_MAGIC):
            return self._legacy_doc(content)
        if ext == ".pdf":
            return self._pdf(content)
        if ext == ".docx":
            return self._docx(content)
        if ext in IMAGE_EXTS:
            return self._image_ocr(content)
        # .txt/.md/.csv/unknown — decode lossily, never raise
        return content.decode("utf-8", errors="ignore").strip()

    # ---- PDF -------------------------------------------------------------
    def _pdf(self, content: bytes) -> str:
        text, needs_ocr = self._pdf_pymupdf4llm(content)
        if text and not needs_ocr:
            return text

        # Fallback to pdfplumber text layer
        pages = self._pdf_text_layer(content)
        text_plumber = "\n".join(pages).strip()

        if not self.ocr:
            return text or text_plumber

        # OCR only if the doc looks scanned: no text layer at all, or some page
        # is near-empty (mixed digital + scanned PDFs are common).
        needs_ocr = needs_ocr or (not pages) or any(len(p.strip()) < _OCR_MIN_CHARS for p in pages)
        if not needs_ocr:
            return text or text_plumber

        ocr_text = self._pdf_ocr(content, pages)
        # prefer the merged OCR result; fall back to whatever text layer existed
        return ocr_text or text or text_plumber

    def _pdf_pymupdf4llm(self, content: bytes) -> tuple[str, bool]:
        """Try pymupdf4llm to extract Markdown. Returns (text, needs_ocr)."""
        try:
            import fitz
            import pymupdf4llm
        except Exception:
            log.warning("pymupdf4llm not installed — cannot extract Markdown from PDF")
            return "", True

        doc = None
        try:
            doc = fitz.Document(stream=content, filetype="pdf")
            text = pymupdf4llm.to_markdown(doc)
            # check if it needs OCR (basically empty)
            needs_ocr = len(text.strip()) < _OCR_MIN_CHARS * len(doc)
            return text, needs_ocr
        except Exception as e:
            log.warning("pymupdf4llm read failed (%s)", e.__class__.__name__)
            return "", True
        finally:
            if doc is not None:
                doc.close()

    @staticmethod
    def _pdf_text_layer(content: bytes) -> list[str]:
        """Per-page text layer. [] if pdfplumber missing or PDF unreadable."""
        try:
            import pdfplumber  # optional dep
        except Exception:
            log.warning("pdfplumber not installed — cannot read PDF text layer")
            return []
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                return [(page.extract_text() or "") for page in pdf.pages]
        except Exception as e:  # encrypted / corrupt / not-a-pdf
            log.warning("PDF text-layer read failed (%s); will try OCR", e.__class__.__name__)
            return []

    def _pdf_ocr(self, content: bytes, pages: list[str]) -> str:
        """Render up to max_ocr_pages and OCR pages lacking a text layer.

        Pages that already have text are kept as-is (no needless OCR). Returns
        "" if OCR deps/binaries are unavailable so the caller can fall back.
        """
        try:
            import pytesseract
            from pdf2image import convert_from_bytes
            from PIL import Image
        except Exception:
            log.warning("OCR skipped: install `pytesseract` and `pdf2image` for scanned PDFs")
            return ""
        Image.MAX_IMAGE_PIXELS = _MAX_IMAGE_PIXELS  # decompression-bomb guard

        # how many pages to consider: the text-layer count if we have it, else the cap.
        # Rendering ALL pages at once (the old convert_from_bytes(last_page=cap)) held
        # every page bitmap in memory simultaneously — ~1.3GB for 50 pages @300DPI.
        # Convert ONE page at a time so peak memory is a single page bitmap.
        limit = min(len(pages), self.max_ocr_pages) if pages else self.max_ocr_pages
        out: list[str] = []
        rendered = 0
        for i in range(1, limit + 1):
            existing = pages[i - 1].strip() if (i - 1) < len(pages) else ""
            if len(existing) >= _OCR_MIN_CHARS:
                out.append(existing)        # has a text layer -> no need to rasterize
                continue
            try:
                imgs = convert_from_bytes(
                    content, dpi=self.ocr_dpi, first_page=i, last_page=i,
                    timeout=self.ocr_timeout_s,
                )
            except Exception as e:
                log.warning("PDF->image failed on page %d (%s); is `poppler` installed?",
                            i, e.__class__.__name__)
                break
            if not imgs:                    # past the last page (unknown-length scan)
                break
            rendered = i
            image = imgs[0]
            try:
                out.append(self._ocr_image(pytesseract, image) or existing)
            finally:
                image.close()               # free the page bitmap before the next one

        merged = "\n".join(p for p in out if p).strip()
        # carry over any text-layer pages beyond what we rendered so nothing is dropped
        if len(pages) > rendered:
            tail = "\n".join(p.strip() for p in pages[rendered:] if p.strip())
            merged = (merged + "\n" + tail).strip() if tail else merged
        return merged

    # ---- image -----------------------------------------------------------
    def _image_ocr(self, content: bytes) -> str:
        if not self.ocr:
            return ""
        try:
            import pytesseract
            from PIL import Image
        except Exception:
            log.warning("OCR skipped: install `pytesseract` and `Pillow` for image files")
            return ""
        Image.MAX_IMAGE_PIXELS = _MAX_IMAGE_PIXELS  # decompression-bomb guard
        try:
            with Image.open(io.BytesIO(content)) as image:
                image.load()
                return self._ocr_image(pytesseract, image)
        except Exception as e:
            # includes Pillow's DecompressionBombError on oversized images
            log.warning("image OCR failed (%s)", e.__class__.__name__)
            return ""

    def _ocr_image(self, pytesseract, image) -> str:
        """One Tesseract call. Returns "" on any engine/binary error."""
        try:
            return pytesseract.image_to_string(
                image, lang=self.ocr_lang, timeout=self.ocr_timeout_s
            ).strip()
        except Exception as e:
            # TesseractNotFoundError, missing lang data, etc.
            log.warning("tesseract OCR error (%s); check tesseract binary + langpack", e.__class__.__name__)
            return ""

    # ---- DOCX ------------------------------------------------------------
    def _docx(self, content: bytes) -> str:
        try:
            import docx  # python-docx, optional dep
        except Exception:
            log.warning("python-docx not installed — cannot read .docx")
            return ""
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as e:
            log.warning("DOCX read failed (%s)", e.__class__.__name__)
            return ""
        parts = [p.text for p in document.paragraphs]
        # include table cells — resumes often put skills/dates in tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        text = "\n".join(t for t in parts if t).strip()
        if self.ocr:
            text = self._docx_image_ocr(document, text)
        return text.strip()

    # ---- legacy .doc (Word 97-2003) --------------------------------------
    def _legacy_doc(self, content: bytes) -> str:
        """Read an OLE2 .doc by shelling out to a converter.

        ponytail: converts rather than parsing the binary Word format — same
        optional-external-binary deal as tesseract/poppler. `textutil` is stock
        on macOS; `soffice` covers Linux/Docker. Tries both because a Homebrew
        `soffice` shim can exist while LibreOffice itself is gone.
        """
        import shutil
        import subprocess
        import tempfile

        with tempfile.TemporaryDirectory() as tmp:
            src = os.path.join(tmp, "in.doc")
            with open(src, "wb") as fh:
                fh.write(content)

            textutil = shutil.which("textutil")          # macOS, always present
            if textutil:
                out = self._run_convert(
                    subprocess, [textutil, "-convert", "txt", "-stdout", src], tmp)
                # textutil exits 0 on a file it can't parse, echoing the raw
                # bytes back. NUL bytes mean it did that — don't let binary
                # sludge reach the screening prompts; try the next converter.
                if out is not None and b"\x00" not in out:
                    return out.decode("utf-8", errors="ignore").strip()

            soffice = shutil.which("soffice") or shutil.which("libreoffice")
            if soffice:
                if self._run_convert(
                    subprocess,
                    [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, src],
                    tmp,
                ) is not None:
                    converted = os.path.join(tmp, "in.docx")
                    if os.path.exists(converted):
                        with open(converted, "rb") as fh:
                            return self._docx(fh.read())

        log.warning(
            "legacy Word .doc (possibly mislabeled .docx) and no working converter — "
            "install libreoffice, or ask for the resume as .docx/PDF"
        )
        return ""

    def _run_convert(self, subprocess, cmd: list[str], tmp: str) -> bytes | None:
        """Run a converter. Returns stdout, or None if it failed/timed out."""
        try:
            done = subprocess.run(
                cmd, capture_output=True, check=True, timeout=self.doc_convert_timeout_s,
                # a stray user profile lock makes soffice hang; give it its own HOME
                env={**os.environ, "HOME": tmp},
            )
        except Exception as e:
            log.warning("%s conversion failed (%s)",
                        os.path.basename(cmd[0]), e.__class__.__name__)
            return None
        return done.stdout

    def _docx_image_ocr(self, document, text: str) -> str:
        """OCR images embedded in a DOCX (some resumes are a single image in a
        Word wrapper). Appends recovered text; logos/icons just yield little.
        Returns `text` unchanged if OCR deps are missing."""
        try:
            import pytesseract
            from PIL import Image
        except Exception:
            return text
        Image.MAX_IMAGE_PIXELS = _MAX_IMAGE_PIXELS  # decompression-bomb guard
        chunks = [text] if text else []
        seen = 0
        for rel in document.part.rels.values():
            if "image" not in rel.reltype or rel.is_external:
                continue
            if seen >= self.max_ocr_pages:
                break
            seen += 1
            try:
                blob = rel.target_part.blob
                with Image.open(io.BytesIO(blob)) as image:
                    image.load()
                    recovered = self._ocr_image(pytesseract, image)
            except Exception as e:
                log.warning("docx image OCR failed (%s)", e.__class__.__name__)
                continue
            if recovered:
                chunks.append(recovered)
        return "\n".join(chunks)


if __name__ == "__main__":  # ponytail: smallest real check of the OCR path
    import sys

    x = FileTextExtractor(ocr_dpi=200)

    # plain text path
    assert x.extract(b"hello world", "a.txt") == "hello world"
    # empty / unknown bytes never raise
    assert x.extract(b"", "a.pdf") == ""

    # content beats extension: a legacy .doc renamed to .docx must NOT hit the
    # zip reader (that was the BadZipFile -> "no readable text" failure), and a
    # truncated OLE2 header must still degrade to "" rather than raise.
    assert x.extract(_OLE2_MAGIC + b"\x00" * 64, "resume.docx") == ""
    # a PDF renamed to .docx routes to the PDF reader
    assert x.extract(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n", "resume.docx") == ""

    # OCR round-trip: render text to a PNG, then read it back via Tesseract.
    try:
        from PIL import Image, ImageDraw, ImageFont
        import pytesseract  # noqa: F401  (ensures binary+lib present before asserting)
    except Exception:
        print("SKIP OCR check: Pillow/pytesseract not installed")
        sys.exit(0)
    try:
        font = ImageFont.load_default(size=48)  # Pillow >=10
    except Exception:
        font = ImageFont.load_default()
    img = Image.new("RGB", (700, 160), "white")
    ImageDraw.Draw(img).text((20, 50), "Resume OCR works", fill="black", font=font)
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    got = x.extract(buf.getvalue(), "scan.png").lower()
    # large clean render -> Tesseract should recover at least one full word
    assert any(w in got for w in ("resume", "works", "ocr")), f"OCR returned: {got!r}"
    print("OK: text, empty, and OCR paths verified")
